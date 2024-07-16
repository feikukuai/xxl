from docx import Document
# 获取当前脚本所在目录的绝对路径
import os
import sys

exe_dir = os.path.dirname(sys.executable)
print(exe_dir)
script_dir = os.path.dirname(sys.executable)
print(script_dir)
# 改变当前工作目录到exe文件所在的目录
os.chdir(exe_dir)
source_dir = os.path.dirname(sys.executable)
print(f"正确工作路径 directory: {os.getcwd()}")
# 需要检查和创建的文件列表
doc_files = ['input.docx','moxing.docx','input1.docx', 'output.docx', 'temp.docx', 'text1.docx','定位编辑.docx','fixtext.docx', 'text2.docx']
# 检查每个文件是否存在，如果不存在则创建一个空文档
for filename in doc_files:
    file_path = os.path.join(script_dir, filename)

    # 如果文件不存在，则创建一个新的空Word文档
    if not os.path.exists(file_path):
        try:
            doc = Document()
            doc.save(file_path)
            print(f"已创建文件: {filename}")
        except Exception as e:
            print(f"无法创建文件{filename}：{e}")
    else:
        print(f"文件 {filename} 已存在，无需操作。")
def read_replacement_rules_from_doc(file_path, delimiter=':'):
    # 打开Word文档
    doc = Document(file_path)

    # 初始化一个空字典存放替换规则
    replacement_rules = {}

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        text = paragraph.text
        parts = text.split(delimiter)
        if len(parts) == 2 and parts[0].strip() and parts[1].strip():
            key = parts[0].strip()
            value = parts[1].strip()
            replacement_rules[key] = value

    return replacement_rules

# 从fixtext.docx中读取替换规则
replacement_rules = read_replacement_rules_from_doc('fixtext.docx')

# 打印读取到的替换规则
print(replacement_rules)

# 检查并打开docx文件
doc_path = os.path.join(script_dir, 'input1.docx')
if not os.path.exists(doc_path):
    raise FileNotFoundError("找不到文件：input1.docx")

doc = Document(doc_path)

# 遍历所有段落进行替换
for paragraph in doc.paragraphs:
    for old_word, new_word in replacement_rules.items():
        paragraph.text = paragraph.text.replace(old_word, new_word)

# 将修改后的内容保存回原文件
doc.save(doc_path)

def remove_empty_paragraphs(file_path):
    # 打开Word文档
    doc = Document(file_path)

    # 遍历所有段落
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if len(paragraph.text.strip()) == 0:  # 检查段落文本是否为空（考虑可能有空白字符）
            paragraphs_to_remove.append(paragraph)

    # 删除空段落（不直接在遍历中删除以避免迭代器错误）
    for paragraph in paragraphs_to_remove:
        paragraph._element.getparent().remove(paragraph._element)

    # 保存修改后的文档
    doc.save(file_path)

# 使用函数处理文件
remove_empty_paragraphs('input1.docx')
# 获取当前脚本所在目录



# 定义替换规则
replacement_rules = {'论文': '洛文','中人': '众人','若问': '洛文','螺纹': '洛文','罗文': '洛文','美丽': '梅莉','微微辣': '薇薇拉','丽雅': '莉娅','需要多': '西奥多',}

# 检查并打开docx文件
doc_path = os.path.join(script_dir, 'input1.docx')
if not os.path.exists(doc_path):
    raise FileNotFoundError("找不到文件：input1.docx")

doc = Document(doc_path)

# 遍历所有段落进行替换
for paragraph in doc.paragraphs:
    for old_word, new_word in replacement_rules.items():
        paragraph.text = paragraph.text.replace(old_word, new_word)

# 将修改后的内容保存回原文件
doc.save(doc_path)
# 现在所有指定的文件都应在当前目录下存在，可以进行后续读取或写入操作。

# 现在所有指定的文件都应在当前目录下存在，可以进行后续读取或写入操作。
import shutil

# 复制文件
shutil.copyfile('output.docx', 'temp.docx')

# 然后使用以下代码清空并保存 'temp.docx' 为 'output.docx'
from docx import Document

# 打开文档
doc1 = Document('temp.docx')

# 清空文档内容并删除空段落
for paragraph in doc1.paragraphs:
    if not paragraph.text.strip():  # 如果段落文本为空或仅包含空格
        paragraph._element.getparent().remove(paragraph._element)  # 删除该段落元素
    else:
        paragraph.clear()

# 保存更改为 'output.docx'
doc1.save('output.docx')