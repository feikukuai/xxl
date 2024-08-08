
import pandas as pd
import openpyxl
import docx
import jieba
from docx import Document
# 获取当前脚本所在目录的绝对路径
import os
import sys

exe_dir = os.path.dirname(sys.executable)
print(exe_dir)
script_dir = os.path.dirname(sys.executable)
print(script_dir)
# 改变当前工作目录到exe文件所在的目录
os.chdir(exe_dir)
current_dir = os.path.dirname(sys.executable)
source_dir = os.path.dirname(sys.executable)
print(f"正确工作路径 directory: {os.getcwd()}")

# 定义读取替换规则的函数
def read_replacement_rules_from_doc(file_path, delimiter=':'):
    doc = Document(file_path)
    replacement_rules = {}
    for paragraph in doc.paragraphs:
        text = paragraph.text
        parts = text.split(delimiter)
        if len(parts) == 2 and parts[0].strip() and parts[1].strip():
            key = parts[0].strip()
            value = parts[1].strip()
            replacement_rules[key] = value
    return replacement_rules

# 从fixtext.docx中读取替换规则
fixtext_path = os.path.join(current_dir, 'input1.docx')
file_path = os.path.join(current_dir, 'input1.docx')
replacement_rules = read_replacement_rules_from_doc(fixtext_path)

docx_file = os.path.join(current_dir, 'input1.docx')

excel_file = os.path.join(current_dir, 'output.xlsx')
# 读取docx文件
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)
import docx
import pandas as pd
from openpyxl import Workbook
from openpyxl import load_workbook
# 读取docx文件
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# 中文分句
def split_sentences(text):
    # 使用中文句号、叹号或问号分割句子
    sentence_delimiters = ['。', '！', '？']
    sentences = []
    temp_sentence = ""
    for char in text:
        temp_sentence += char
        if char in sentence_delimiters:
            sentences.append(temp_sentence.strip())
            temp_sentence = ""
    return sentences

# 存储到现有的Excel，根据数字决定工作表
def save_to_excel(sentences, excel_path, sheet_index):
    wb = load_workbook(excel_path)
    ws = wb.worksheets[sheet_index - 1]  # 工作表索引从0开始
    row = 8  # Excel是从1开始的，所以第8行是7

    for sentence in sentences:
        ws.cell(row=row, column=2, value=sentence)  # B列是第2列
        row += 1
    wb.save(excel_path)

# 主函数
def main():
   
    text = read_docx(docx_file)
    
    # 提取第一个数字作为工作表索引
    first_line = text.split('\n', 1)[0]  # 获取第一行
    sheet_index = int(first_line)  # 将第一行转换为整数作为工作表索引
    
    # 移除第一行
    text = text[len(first_line):].strip()
    
    sentences = split_sentences(text)
    save_to_excel(sentences, excel_file, sheet_index)
    print(f"句子已保存到Excel的第{sheet_index}个工作表中，从B8开始。")

if __name__ == "__main__":
    main()
