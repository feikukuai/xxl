from docx import Document
# 获取当前脚本所在目录的绝对路径
import os
import configparser
# 获取Python解释器（或exe）所在目录
exe_dir = os.path.dirname(sys.executable)
print(exe_dir)
script_dir = os.path.dirname(sys.executable)
print(script_dir)
# 改变当前工作目录到exe文件所在的目录
os.chdir(exe_dir)
source_dir = os.path.dirname(sys.executable)
print(f"正确工作路径 directory: {os.getcwd()}")
# 需要检查和创建的文件列表
doc_files = ['input.docx','moxing.docx','input1.docx', 'output.docx', 'temp.docx', 'text3.docx','定位编辑.docx','fixtext.docx', 'text2.docx']
# 检查每个文件是否存在，如果不存在则创建一个空文档
for filename in doc_files:
    file_path = os.path.join(script_dir, filename)

    # 如果文件不存在，则创建一个新的空Word文档
    if not os.path.exists(file_path):
        try:
            doc = Document()
            doc.save(file_path)
            print(f"已创建文件: {filename}")
        except Exception as e:
            print(f"无法创建文件{filename}：{e}")
    else:
        print(f"文件 {filename} 已存在，无需操作。")
def read_replacement_rules_from_doc(file_path, delimiter=':'):
    # 打开Word文档
    doc = Document(file_path)

    # 初始化一个空字典存放替换规则
    replacement_rules = {}

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        text = paragraph.text
        parts = text.split(delimiter)
        if len(parts) == 2 and parts[0].strip() and parts[1].strip():
            key = parts[0].strip()
            value = parts[1].strip()
            replacement_rules[key] = value

    return replacement_rules

# 从fixtext.docx中读取替换规则
replacement_rules = read_replacement_rules_from_doc('fixtext.docx')

# 打印读取到的替换规则
print(replacement_rules)

# 检查并打开docx文件
doc_path = os.path.join(script_dir, 'input1.docx')
input_file_path = doc_path
if not os.path.exists(doc_path):
    raise FileNotFoundError("找不到文件：input1.docx")

doc = Document(doc_path)

# 遍历所有段落进行替换
for paragraph in doc.paragraphs:
    for old_word, new_word in replacement_rules.items():
        paragraph.text = paragraph.text.replace(old_word, new_word)

# 将修改后的内容保存回原文件
doc.save(doc_path)

def remove_empty_paragraphs(file_path):
    # 打开Word文档
    doc = Document(file_path)

    # 遍历所有段落
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if len(paragraph.text.strip()) == 0:  # 检查段落文本是否为空（考虑可能有空白字符）
            paragraphs_to_remove.append(paragraph)

    # 删除空段落（不直接在遍历中删除以避免迭代器错误）
    for paragraph in paragraphs_to_remove:
        paragraph._element.getparent().remove(paragraph._element)

    # 保存修改后的文档
    doc.save(file_path)

# 使用函数处理文件
remove_empty_paragraphs('input1.docx')
# 获取当前脚本所在目录



# 定义替换规则
replacement_rules = {'论文': '洛文',}

# 检查并打开docx文件
doc_path = os.path.join(script_dir, 'input1.docx')
if not os.path.exists(doc_path):
    raise FileNotFoundError("找不到文件：input1.docx")

doc = Document(doc_path)

# 遍历所有段落进行替换
for paragraph in doc.paragraphs:
    for old_word, new_word in replacement_rules.items():
        paragraph.text = paragraph.text.replace(old_word, new_word)

# 将修改后的内容保存回原文件
doc.save(doc_path)
# 现在所有指定的文件都应在当前目录下存在，可以进行后续读取或写入操作。

# 现在所有指定的文件都应在当前目录下存在，可以进行后续读取或写入操作。
import shutil

# 复制文件
shutil.copyfile('output.docx', 'temp.docx')

# 然后使用以下代码清空并保存 'temp.docx' 为 'output.docx'
from docx import Document

# 打开文档
doc1 = Document('temp.docx')

# 清空文档内容并删除空段落
for paragraph in doc1.paragraphs:
    if not paragraph.text.strip():  # 如果段落文本为空或仅包含空格
        paragraph._element.getparent().remove(paragraph._element)  # 删除该段落元素
    else:
        paragraph.clear()

# 保存更改为 'output.docx'
doc1.save('output.docx')

os.environ['DASHSCOPE_API_KEY'] = 'sk-ade26912d9f6406fabe8edd7c5b2b7b1'
from http import HTTPStatus
import dashscope
import json
from http import HTTPStatus
import dashscope
from docx import Document


def merge_two_docs(doc1_path, doc2):
    # 读取第一个文档
    doc1 = Document(doc1_path)

    # 将第二个文档的内容添加到第一个文档的末尾
    for paragraph in doc2.paragraphs:
        doc1.add_paragraph(paragraph.text)

    # 保存合并后的文档
    doc1.save(doc1_path)


def save_to_docx(response):
    doc = Document()

    for choice in response.output['choices']:
        content = choice['message']['content']
        content += 'A'  # 在内容末尾添加大写的'A'
        doc.add_paragraph(content)

    # 合并文档
    merge_two_docs('output.docx', doc)
    print("内容已添加到output.docx")


def read_text_from_doc(file_path, batch_size=500, min_batch_size=500, setup_info="text2 "):
    doc = Document(file_path)
    text_batches = []
    current_batch = ""
    text2 = ""
    for paragraph in doc.paragraphs:
        current_paragraph_text = paragraph.text.strip()
        if len(current_batch + current_paragraph_text) <= batch_size:
            current_batch += f"{setup_info}{current_paragraph_text}\n"

        else:
            # 如果当前段落使总字符数超过batch_size，则将当前批次添加到text_batches并开始新的批次
            text_batches.append(current_batch)
            current_batch = f"{setup_info}{current_paragraph_text}\n"


    # 处理最后一批次，如果小于min_batch_size，则将其与前一个批次合并
    if len(current_batch) >= min_batch_size:

        text_batches.append(current_batch)
    elif len(text_batches) > 0:
        text_batches[-1] += current_batch
    # 获取用户输入




    user_input = None

    user_input = 1
    print("你选择了数字1")


    # 读取文档text3和text2
    def read_docx(filename):
        document = Document(filename)
        text_content = [paragraph.text for paragraph in document.paragraphs]
        return '\n'.join(text_content)

    # 获取当前脚本所在目录

    # 计算文本文件的相对路径
    file_path_text3 = os.path.join(script_dir, 'text3.docx')
    file_path_text2 = os.path.join(script_dir, 'text2.docx')

    text3 = read_docx(file_path_text3)
    text2 = read_docx(file_path_text2)

    print("text3的内容：")
    print(text3)

    print("\ntext2的内容：")
    print(text2)

    text0 = ""
    # 根据用户输入给text0赋值
    if user_input == 1:
        text2 = ""
    elif user_input == 2:
        text3 = ""
    else:
        # 如果输入既不是1也不是2，则可以给出错误提示或其他处理方式
        print("无效输入，请输入1或2")

    for i in range(len(text_batches)):
        text_batches[i] = text3 + "" + text_batches[i] + "" + text2

    # for i in range(len(text_batches)):
    #  text_batches[i] = text_batches[i] + "" + text2
    return text_batches

config = configparser.ConfigParser()

# 读取配置文件
config.read('config.ini', encoding='utf-8')
digit = config.get('参量配置', 'digit')
digit = int(digit)
ConfigurationNumber = config.get('参量配置', 'ConfigurationNumber')
ConfigurationNumber = float(ConfigurationNumber)
choose = config.get('参量配置', 'choose')
section = f'数据库配置{choose}'
company = config.get(section, 'company')
model = config.get(section, 'model')
api = config.get(section, 'api')
temperature = config.get(section, 'temperature')
temperature = float(temperature)

print(f"Company: {company}")
print(f"Model: {model}")
print(f"API: {api}")
print(f"Temperature: {temperature}")

# 确保ai处理的函数独立
from openai import OpenAI



def add_newline_after_comma(docx_path, output_path):
    doc = Document(docx_path)
    for para in doc.paragraphs:
        text = para.text
        new_text = []
        i = 0
        while i < len(text):
            if text[i] == '”':
                # 查找前面最近的“
                j = i - 1
                while j >= 0 and text[j] != '“':
                    j -= 1
                # 如果找到了“，并且距离超过8个字符
                if j >= 0 and (i - j) > 8:
                    new_text.append('”\n')
                else:
                    new_text.append('”')
            else:
                new_text.append(text[i])
            i += 1
        para.text = ''.join(new_text)
    doc.save(output_path)

# 使用示例
# add_newline_after_comma('input.docx', 'output.docx')

def remove_empty_paragraphs(doc_path):
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    for para in paragraphs:
        if not para.text.strip():
            p = para._element
            p.getparent().remove(p)
    doc.save(doc_path)



def culi(a, api_key,fieldQ):
    # 创建 OpenAI 客户端实例
    client = OpenAI(api_key=api, base_url=company)

    # 使用传入的消息列表a进行聊天
    response = client.chat.completions.create(
        model=model,
        messages=a, 
        temperature=temperature 
    )

    # 将API的响应添加到消息列表中

    a = response.choices[0].message
    result_string = str(a)
    content_start = result_string.find('content=')
    if content_start != -1:
        content_start += len('content=')
        content_end = result_string.find(', refusal=None', content_start)
        if content_end == -1:
            content_end = len(result_string)
        content = result_string[content_start:content_end].strip('"')


    sd_content = content.strip("'")
    sd_content = sd_content.replace('\\n', '\n')
    sd_content = sd_content + "\nA"
    print(sd_content)
    doc = Document('output.docx')
    chinese_punctuation = "，。……！？…………；：、（）〈〉《》{}【】“”‘’"
    from fuzzywuzzy import process
    # 获取所有可能的匹配项，按分数排序
    import re

    sentences = re.split(r'([' + re.escape(chinese_punctuation) + '])', sd_content)
    # 将标点符号重新拼接到句子上
    sentences = [sentences[i] + sentences[i+1] for i in range(0, len(sentences)-1, 2)]
    pattern = re.compile(r".*[" + re.escape(chinese_punctuation) + "]$")
    candidates = [s for s in sentences if pattern.match(s)]
    matches = process.extract(fieldQ, candidates, limit=20)

# 遍历匹配项，找到 matched_word
    matched_word = None
    score = 0
    for match, match_score in matches:
        if match and match[-1] in chinese_punctuation and match_score > pipeisuzi:
           matched_word = match
           score = match_score
           break

    if matched_word:
       position = sd_content.find(matched_word)
       if position != -1:
           sd_content = sd_content[:position + len(matched_word)]
    else:
        sd_content = sd_content
    doc.add_paragraph(sd_content)

    # 保存修改后的文档
    doc.save('output.docx')
    add_newline_after_comma('output.docx', 'output.docx')
    remove_empty_paragraphs('output.docx')


api_key = api  # 应该从安全的地方获取
suzi = digit
suzi = int(suzi)
pipeisuzi = ConfigurationNumber
if __name__ == '__main__':
    text2 = ""
    text_batches = read_text_from_doc(input_file_path, setup_info=text2)
    for i, text_batch in enumerate(text_batches):
        fieldQ = text_batch[-suzi:]
        messages = [{"role": "user", "content": text_batch}]
        culi(messages, api_key,fieldQ)
from docx import Document

# 使用您的实际文件路径替换 'your_file_path.docx'

# 获取当前脚本所在目录的绝对路径


# 构建output.docx文件的绝对路径
file_path = os.path.join(script_dir, 'output.docx')
print(file_path)
print("位置在于:")


print(source_dir)

remove_empty_paragraphs('output.docx')
# 使用示例
# add_newline_after_comma('input.docx', 'output.docx')
# 调用函数，替换成你的文档路径
